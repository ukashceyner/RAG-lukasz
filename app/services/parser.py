"""Document parsing service for PDF and DOCX extraction."""

import io
import logging
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document

logger = logging.getLogger(__name__)


class DocumentParser:
    """Handles text extraction from PDF and DOCX files."""

    @staticmethod
    async def parse(file_content: bytes, filename: str) -> str:
        """
        Extract text from a document file.

        Args:
            file_content: Raw bytes of the uploaded file
            filename: Original filename to determine file type

        Returns:
            Extracted text content as a string

        Raises:
            ValueError: If file type is not supported
        """
        extension = Path(filename).suffix.lower()

        if extension == ".pdf":
            return await DocumentParser._parse_pdf(file_content)
        elif extension == ".docx":
            return await DocumentParser._parse_docx(file_content)
        else:
            raise ValueError(f"Unsupported file type: {extension}")

    @staticmethod
    async def _parse_pdf(file_content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            text_parts = []
            with fitz.open(stream=file_content, filetype="pdf") as doc:
                for page_num, page in enumerate(doc):
                    page_text = page.get_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                    logger.debug(f"Extracted text from PDF page {page_num + 1}")

            full_text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from PDF")
            return full_text

        except Exception as e:
            logger.error(f"Failed to parse PDF: {e}")
            raise ValueError(f"Failed to parse PDF: {e}")

    @staticmethod
    async def _parse_docx(file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(io.BytesIO(file_content))
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            full_text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from DOCX")
            return full_text

        except Exception as e:
            logger.error(f"Failed to parse DOCX: {e}")
            raise ValueError(f"Failed to parse DOCX: {e}")


# Singleton instance
parser = DocumentParser()
